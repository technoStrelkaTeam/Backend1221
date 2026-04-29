from pathlib import Path
from typing import List
from migration.readers.base import DocumentReader, DocumentChunk


class WordDocumentReader(DocumentReader):
    extension = ".docx"

    extensions = {".docx", ".doc"}

    async def read(self, file_path: Path) -> List[DocumentChunk]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("pip install python-docx")

        def _read_sync(path):
            chunks = []

            doc = Document(path)

            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            if full_text:
                chunks.append(DocumentChunk(
                    key=self.get_key(file_path.stem),
                    title=file_path.stem,
                    text="\n".join(full_text)
                ))

            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        table_text.append(" | ".join(row_cells))

                if table_text:
                    chunks.append(DocumentChunk(
                        key=f"{file_path.stem}_table",
                        title=f"{file_path.stem} - Таблица",
                        text="\n".join(table_text)
                    ))

            if not chunks:
                chunks.append(DocumentChunk(
                    key=self.get_key(file_path.stem),
                    title=file_path.stem,
                    text="Empty Word document"
                ))

            return chunks

        return await self._run_sync(_read_sync, file_path)